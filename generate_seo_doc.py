"""Generate BCW Homepage SEO Strategy document as .docx"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

TEAL = RGBColor(0x13, 0x53, 0x50)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TEAL if level <= 2 else DARK
    return h


def add_quote(text):
    """Add a blockquote-style paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'\u201c{text}\u201d')
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


def add_source(name, url):
    """Add a source line."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'Bron: {name}')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.font.bold = True
    p.add_run('\n')
    link_run = p.add_run(url)
    link_run.font.size = Pt(8.5)
    link_run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
    return p


def add_label(label, text):
    """Add a bold label followed by normal text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run_label = p.add_run(f'{label}: ')
    run_label.font.bold = True
    run_label.font.color.rgb = DARK
    run_text = p.add_run(text)
    run_text.font.color.rgb = DARK
    return p


# ============================================================
# DOCUMENT CONTENT
# ============================================================

# -- Title --
title = doc.add_heading('Beton Cire Webshop', level=0)
for run in title.runs:
    run.font.color.rgb = TEAL
    run.font.size = Pt(26)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = subtitle.add_run('Homepage SEO Verbetervoorstel')
run.font.size = Pt(16)
run.font.color.rgb = GRAY
run.font.bold = True

doc.add_paragraph('')

# Status box
status = doc.add_paragraph()
status.paragraph_format.space_before = Pt(8)
status.paragraph_format.space_after = Pt(8)
run = status.add_run('Status: ')
run.font.bold = True
run.font.size = Pt(10)
run = status.add_run('Voorstel ter review \u2014 nog niet doorgevoerd op productie. ')
run.font.size = Pt(10)
run = status.add_run('Wijzigingen staan op staging ter beoordeling.')
run.font.size = Pt(10)

doc.add_paragraph('\u2500' * 60)

# ============================================================
# 1. ALT TEXTS
# ============================================================
add_heading_styled('1. Ontbrekende alt-teksten op afbeeldingen', level=1)

add_label('Probleem',
          '5 afbeeldingen op de homepage hebben geen alt-tekst (leeg alt-attribuut). '
          'Dit betreft de hero banner, kleurstalen CTA, All-In-One productafbeelding, '
          'Original productafbeelding, en de inspiratie-sectie afbeelding.')

add_label('Voorgestelde fix',
          'Beschrijvende, beknopte alt-teksten toevoegen die de afbeelding in context plaatsen.')

p = doc.add_paragraph()
run = p.add_run('Bronnen:')
run.font.bold = True

add_quote('When writing alt text, focus on creating useful, information-rich content '
          'that uses keywords appropriately and is in context of the content of the page.')
add_quote('Avoid filling alt attributes with keywords (also known as keyword stuffing) '
          'as it results in a negative user experience and may cause your site to be seen as spam.')
add_source('Google Search Central \u2014 Image SEO Best Practices',
           'https://developers.google.com/search/docs/appearance/google-images')

add_quote('Writing good alt text is quite important. A short, but descriptive piece of text '
          'that explains the relationship between the image and your content.')
add_quote('Google uses alt text along with computer vision algorithms and the contents of the '
          'page to understand the subject matter of the image.')
add_source('Google SEO Starter Guide',
           'https://developers.google.com/search/docs/fundamentals/seo-starter-guide')

# ============================================================
# 2. OPEN GRAPH
# ============================================================
add_heading_styled('2. Open Graph tags ontbreken voor de homepage', level=1)

add_label('Probleem',
          'De homepage heeft geen og:title, og:description, of og:image ingesteld '
          '(niet in Yoast Social instellingen, niet per pagina). Wanneer iemand de homepage '
          'deelt op Facebook, LinkedIn of Twitter/X verschijnt een generieke of willekeurige preview.')

add_label('Voorgestelde fix',
          'In Yoast SEO > Social de frontpage OG title, description en image instellen.')

p = doc.add_paragraph()
run = p.add_run('Bronnen:')
run.font.bold = True

add_quote('Without these Open Graph tags, the Facebook Crawler uses internal heuristics to make '
          'a best guess about the title, description, and preview image for your content. '
          'Designate this info explicitly with Open Graph tags to ensure the highest quality posts on Facebook.')
add_source('Meta/Facebook \u2014 Sharing Webmasters Documentation',
           'https://developers.facebook.com/docs/sharing/webmasters/')

add_quote('Maximize the click through rates of your stories by specifying Open Graph tags '
          'for all your articles and including compelling images, titles and descriptions.')
add_source('Meta/Facebook \u2014 Best Practices',
           'https://developers.facebook.com/docs/sharing/best-practices')

add_quote('Use images that are at least 1200 x 630 pixels for the best display on high resolution devices. '
          'The minimum allowed image dimension is 200 x 200 pixels. '
          'Try to keep your images as close to 1.91:1 aspect ratio as possible.')
add_source('Meta/Facebook \u2014 Image Requirements',
           'https://developers.facebook.com/docs/sharing/webmasters/images/')

add_source('Open Graph Protocol Specificatie (og:title, og:type, og:image, og:url zijn vereist)',
           'https://ogp.me/')

add_source('Yoast \u2014 Open Graph configuratie',
           'https://yoast.com/help/custom-open-graph-tags/')

# ============================================================
# 3. PAGE TITLE
# ============================================================
add_heading_styled('3. Paginatitel is "HOME" (niet beschrijvend)', level=1)

add_label('Probleem',
          'De WordPress post_title van de homepage is "HOME". Hoewel Yoast de <title> tag overschrijft, '
          'kan deze generieke titel doorlekken naar breadcrumbs, RSS feeds, interne zoekresultaten, '
          'en Google kan deze als fallback gebruiken voor title links.')

add_label('Voorgestelde fix',
          'Wijzigen naar "Beton Cire Webshop" of een andere beschrijvende titel.')

p = doc.add_paragraph()
run = p.add_run('Bronnen:')
run.font.bold = True

add_quote('Avoid vague descriptors like "Home" for your home page, '
          'or "Profile" for a specific person\'s profile.')
add_quote('Write descriptive and concise text for your <title> elements.')
add_quote('Title links are critical to giving users a quick insight into the content of a result '
          'and why it\'s relevant to their query. It\'s often the primary piece of information '
          'people use to decide which result to click.')
add_source('Google Search Central \u2014 Title Links',
           'https://developers.google.com/search/docs/appearance/title-link')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run('Let op: ')
run.font.bold = True
run = p.add_run('Google genereert title links uit meerdere bronnen: '
                '<title> element, H1 headings, og:title, en andere prominente tekst op de pagina.')

# ============================================================
# 4. PRICING IN HEADINGS
# ============================================================
add_heading_styled('4. Prijsinformatie in heading tags (H3)', level=1)

add_label('Probleem',
          '"Vanaf", "\u20ac28", "per 1m\u00b2" (en vergelijkbare prijzen) staan elk in een apart <h3> tag. '
          'Heading tags zijn bedoeld voor content-structuur, niet voor prijsweergave of visuele opmaak.')

add_label('Voorgestelde fix',
          'Wijzigen van <h3> naar <p> tags met CSS-klassen voor visuele opmaak.')

p = doc.add_paragraph()
run = p.add_run('Bronnen:')
run.font.bold = True

add_quote('Don\'t use headings to make text bold or large; use CSS instead.')
add_quote('It is recommended to use heading levels similarly to heading levels in a text editor: '
          'starting with a <h1> as the main heading, with <h2> as headings for sub-sections, '
          'and <h3> if those sub-sections have sections; avoid skipping heading levels.')
add_source('web.dev \u2014 Headings and Sections',
           'https://web.dev/learn/html/headings-and-sections')

add_quote('Nest headings by their rank (or level). The most important heading has the rank 1 (<h1>), '
          'the least important heading rank 6 (<h6>).')
add_source('W3C WAI \u2014 Page Structure Headings',
           'https://www.w3.org/WAI/tutorials/page-structure/headings/')

# ============================================================
# 5. DUPLICATE H1/H2
# ============================================================
add_heading_styled('5. Dubbele H1/H2 tekst', level=1)

add_label('Probleem',
          'De H1 tekst "Een unieke stijl met Beton Cir\u00e9: D\u00e9 trend voor vloeren en wanden" '
          'wordt verderop op de pagina herhaald als H2. Identieke heading-tekst op dezelfde pagina '
          'maakt de structuur onduidelijk.')

add_label('Voorgestelde fix',
          'De dubbele H2 een andere, unieke tekst geven. '
          'De exacte tekst is ter beoordeling door de SEO specialist.')

p = doc.add_paragraph()
run = p.add_run('Bronnen:')
run.font.bold = True

add_quote('Ensure each page has a unique h1 heading.')
add_quote('It\'s easier to jump between pages and sections of a page '
          'if the headings and titles are unique.')
add_source('Google Developer Documentation Style Guide \u2014 Headings',
           'https://developers.google.com/style/headings')

# ============================================================
# 6. HEADING LEVEL SKIPPING
# ============================================================
add_heading_styled('6. Heading levels overgeslagen (H1 \u2192 H4)', level=1)

add_label('Probleem (staging)',
          '"Zeker zijn van je kleur?" was gewijzigd van H2 (productie) naar H4, '
          'wat heading levels H2 en H3 overslaat. Evenzo was "Beton cire vloer All in one" '
          'van H3 naar H4 gewijzigd.')

add_label('Voorgestelde fix',
          'Terugzetten naar H2 respectievelijk H3 (de originele productie-waarden).')

p = doc.add_paragraph()
run = p.add_run('Bronnen:')
run.font.bold = True

add_quote('Skipping heading ranks can be confusing and should be avoided where possible: '
          'Make sure that a <h2> is not followed directly by an <h4>.')
add_source('W3C WAI \u2014 Page Structure Headings',
           'https://www.w3.org/WAI/tutorials/page-structure/headings/')

add_quote('Do not skip heading levels: always start from <h1>, followed by <h2> and so on.')
add_source('MDN Web Docs \u2014 Heading Elements',
           'https://developer.mozilla.org/en-US/docs/Web/HTML/Reference/Elements/Heading_Elements')

add_quote('Headings must be in a valid logical order, meaning h1 through h6 element tags '
          'must appear in a sequentially-descending order. '
          'For example, the heading level following an h1 element should be an h2 element, '
          'not an h3 element.')
add_source('Deque University \u2014 Axe Accessibility Rules',
           'https://dequeuniversity.com/rules/axe/4.4/heading-order')

add_quote('Some screen reader users navigate by headings to understand page structure, '
          'making proper heading hierarchy essential for accessibility.')
add_source('web.dev \u2014 Headings and Sections',
           'https://web.dev/learn/html/headings-and-sections')

# ============================================================
# 7. BELANGRIJKSTE PUNTEN
# ============================================================
add_heading_styled('7. "Belangrijkste punten" van H2 naar H3', level=1)

add_label('Probleem',
          '"Belangrijkste punten" stond als H2 maar is een subsectie van de sectie erboven. '
          'Het hoort een H3 te zijn in de hierarchie.')

add_label('Bronnen', 'Dezelfde als punt 4 en 6 hierboven \u2014 headings moeten de '
          'content-hierarchie weerspiegelen.')

# ============================================================
# 8. DATA-PM-SLICE
# ============================================================
add_heading_styled('8. data-pm-slice attributen opschonen', level=1)

add_label('Probleem',
          'Headings bevatten data-pm-slice="1 1 []" attributen. '
          'Dit zijn ProseMirror editor-artefacten die geen semantische waarde hebben.')

add_label('Voorgestelde fix',
          'Verwijderen. Deze attributen zijn onzichtbaar voor gebruikers maar maken de HTML onnodig romelig.')

# ============================================================
# 9. H4 NBSP
# ============================================================
add_heading_styled('9. H4 captions met overmatige spaties', level=1)

add_label('Probleem',
          'De H4 "In het toilet" bevat 17 non-breaking spaces voor de tekst als visuele positionering. '
          'Dit is geen valide manier om elementen te positioneren.')

add_label('Voorgestelde fix',
          'Spaties verwijderen, CSS gebruiken voor positionering.')

# ============================================================
# OVERIGE BEVINDINGEN
# ============================================================
doc.add_page_break()
add_heading_styled('Overige bevindingen (niet gewijzigd, ter beoordeling)', level=1)

add_heading_styled('A. Alle content nog in Flatsome shortcodes', level=2)
p = doc.add_paragraph(
    'De homepage-content bestaat nog uit [ux_banner], [row], [col] shortcodes. '
    'De overhaul naar Gutenberg blocks is nog nodig voor schonere, semantische HTML.')

add_heading_styled('B. Twitter/X site handle ontbreekt', level=2)
p = doc.add_paragraph(
    'twitter_site staat leeg in Yoast Social instellingen. '
    'Als er een X/Twitter account is, kan deze worden ingevuld.')
add_source('X Developer Platform \u2014 Cards',
           'https://developer.x.com/en/docs/x-for-websites/cards/guides/getting-started')

add_heading_styled('C. Staging go-live checklist', level=2)
items = [
    'blog_public van 0 naar 1 zetten',
    'Fysieke robots.txt verwijderen',
    'X-Robots-Tag header uit .htaccess verwijderen',
    'Database search-replace: staging.beton-cire-webshop.nl \u2192 beton-cire-webshop.nl',
    'Ontbrekende plugins installeren (GTM, reviews, product feed, cookie consent)',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

# ============================================================
# NUANCE SECTION
# ============================================================
doc.add_paragraph('')
doc.add_paragraph('\u2500' * 60)

add_heading_styled('Nuance: Google over heading-structuur en ranking', level=1)

p = doc.add_paragraph()
run = p.add_run('Google\'s SEO Starter Guide zegt het volgende over headings en ranking:')

add_quote('Having your headings in semantic order is fantastic for screen readers.')
add_quote('There\'s also no magical, ideal amount of headings a given page should have.')
add_quote('From Google Search perspective, it doesn\'t matter if you\'re using them out of order.')

p = doc.add_paragraph()
run = p.add_run('Dit betekent dat heading-structuur geen directe ranking factor is, '
                'maar wel invloed heeft op:')

items = [
    'Toegankelijkheid (WCAG/screenreaders)',
    'Gebruikerservaring',
    'Hoe Google de content-structuur interpreteert voor featured snippets en title link generatie',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_source('Google SEO Starter Guide',
           'https://developers.google.com/search/docs/fundamentals/seo-starter-guide')

# ============================================================
# SAVE
# ============================================================
out_path = os.path.join(os.path.dirname(__file__), 'BCW-Homepage-SEO-Strategie.docx')
doc.save(out_path)
print(f'Document saved to: {out_path}')
