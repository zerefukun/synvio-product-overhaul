"""Generate DOCX document: smartdeco.nl migration preparation summary."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Title ---
title = doc.add_heading('Smartdeco.nl - Migratie-voorbereiding', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Uitgevoerd op 10 april 2026')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# --- Inleiding ---
doc.add_heading('Samenvatting', level=1)
doc.add_paragraph(
    'Op 10 april 2026 is smartdeco.nl voorbereid op de migratie naar epoxystone-gietvloer.nl. '
    'Alle oude content die niet meegaat is gemarkeerd voor deindexatie (410 Gone), '
    'de webshop is uitgeschakeld, en de site is achter een onderhoudspagina geplaatst. '
    'URLs die epoxystone straks gaat gebruiken zijn bewust intact gelaten.'
)

# =============================================
# 1. Under Construction
# =============================================
doc.add_heading('1. Under Construction Page ingeschakeld', level=1)
doc.add_paragraph(
    'De plugin "Under Construction Page" is geactiveerd. Bezoekers zien nu een onderhoudspagina '
    'met de tekst "We zijn bezig met werkzaamheden op de site". Alleen ingelogde administrators '
    'kunnen de echte site zien.'
)
doc.add_paragraph('Effect op SEO:', style='List Bullet')
items = [
    'Google krijgt een 503 Service Unavailable response',
    'Dit vertelt Google: "de site is tijdelijk offline, kom later terug"',
    'Google behoudt de huidige index maar stopt met crawlen',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet 2')

# =============================================
# 2. WooCommerce uitgeschakeld
# =============================================
doc.add_heading('2. WooCommerce en betaal-plugins uitgeschakeld', level=1)
doc.add_paragraph(
    'WooCommerce en alle gerelateerde plugins zijn gedeactiveerd. '
    'Er kunnen geen bestellingen meer binnenkomen op het oude smartdeco.'
)

table = doc.add_table(rows=1, cols=2, style='Light List Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Plugin'
hdr[1].text = 'Status'

plugins = [
    ('WooCommerce', 'Inactive'),
    ('Mollie Payments for WooCommerce', 'Inactive'),
    ('WooCommerce PDF Invoices & Packing Slips', 'Inactive'),
    ('WooCommerce Services', 'Inactive'),
    ('Woo Parcel Pro', 'Inactive'),
    ('SendCloud Shipping', 'Inactive'),
    ('YITH WooCommerce Product Add-Ons', 'Inactive'),
    ('YITH WooCommerce Advanced Product Options', 'Inactive'),
    ('YITH WooCommerce Wishlist', 'Inactive'),
]
for name, status in plugins:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = status

doc.add_paragraph()

# =============================================
# 3. Zoekmachines ontmoedigen
# =============================================
doc.add_heading('3. Zoekmachines ontmoedigen (blog_public = 0)', level=1)
doc.add_paragraph(
    'De WordPress instelling "Discourage search engines from indexing this site" staat aan. '
    'Dit voegt het volgende toe aan alle pagina\'s:'
)
doc.add_paragraph('<meta name="robots" content="noindex, nofollow">', style='List Bullet')
doc.add_paragraph('Disallow: / in de virtuele robots.txt', style='List Bullet')
doc.add_paragraph(
    'Dit is een extra veiligheidslaag bovenop de 410-regels en de onderhoudspagina.'
)

# =============================================
# 4. 410 Gone regels
# =============================================
doc.add_heading('4. 410 Gone regels in .htaccess', level=1)
doc.add_paragraph(
    'Een backup is gemaakt als .htaccess.bak-20260410. Daarna zijn 410 Gone regels toegevoegd. '
    'HTTP 410 vertelt zoekmachines: "deze URL bestond ooit maar is permanent verwijderd, '
    'stop met crawlen en deindexeer deze URL."'
)
doc.add_paragraph(
    'Dit is effectiever dan noindex of 404 omdat Google een 410 sneller uit de index haalt.'
)

# -- Table: 410 regels --
doc.add_heading('Overzicht geblokkeerde URLs', level=2)

table2 = doc.add_table(rows=1, cols=3, style='Light List Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Categorie'
hdr2[1].text = '.htaccess regel'
hdr2[2].text = 'Voorbeelden'

gone_rules = [
    ('Blogposts (datum-gebaseerd)',
     '^20[0-9]{2}/[0-9]{2}/[0-9]{2}/',
     '/2024/02/15/blog-titel/'),
    ('Beton-cire producten',
     '^product/beton-cire-',
     '/product/beton-cire-sandy-beach, /product/beton-cire-pearl-white, etc.'),
    ('Kant-klare cementpasta producten',
     '^product/microcement-kant-klare-cementpasta-',
     '/product/microcement-kant-klare-cementpasta-cement-1, -sand-1, etc.'),
    ('Flatsome demo-content',
     '^featured_item, ^featured_item_category',
     '/featured_item/xyz/'),
    ('Oude pagina-slugs',
     'Individuele regels per slug',
     'beton-cire-kleuren, beton-cire-showroom, beton-cire-toepassingen, '
     'beton-cire-producten, beton-cire-ontdekken, handleiding-beton-cire-all-in-one, '
     'handleiding-beton-cire, technische-pagina'),
    ('Kleurcategorieen',
     '^product-category/(beige|blauw|bruin|...)',
     '/product-category/beige, /product-category/grijs, etc. '
     '(beige, blauw, bruin, grijs, groen, roze, rood, taupe, wit, zand, zwart)'),
    ('Oude categorieen',
     'Individuele regels',
     '/product-category/gereedschap, /product-category/losse-materialen, '
     '/product-category/uncategorized'),
    ('Smartdeco-only producten',
     'Individuele regels per product',
     'gereedschapsset, vachtroller-25cm, flexibele-spaan-zonder-handvat, '
     'kim-en-hoek-afdichtingset, verfbak, troffel-180mm, presealer-5m*, '
     'primer-5m2, pu-watervast'),
    ('Smartdeco-only pagina\'s',
     'Individuele regels per pagina',
     'tips-and-tricks, technische, toepassingen, showroom, '
     'handleiding-microcement, handleiding-epoxystone, beton-cire-webshop'),
    ('Oude shop-pagina',
     '^producten',
     '/producten/, /producten/page/2/'),
    ('Author pages',
     '^author/',
     '/author/admin/'),
    ('Ongeldige tel: links',
     '/tel:',
     'URLs met /tel: die in Google stonden'),
]

for cat, rule, examples in gone_rules:
    row = table2.add_row().cells
    row[0].text = cat
    row[1].text = rule
    row[2].text = examples

# Set column widths
for row in table2.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(2.2)
    row.cells[2].width = Inches(3.0)

doc.add_paragraph()

# =============================================
# 5. Veilige URLs
# =============================================
doc.add_heading('5. URLs die NIET geblokkeerd zijn (bewaard voor epoxystone)', level=1)
doc.add_paragraph(
    'De volgende URLs zijn bewust intact gelaten. Deze worden straks door epoxystone '
    'gebruikt op smartdeco.nl na de migratie:'
)

safe_categories = [
    ('Pagina\'s', '/contact, /microcement, /microcement-kleuren, /kleurstalen, '
     '/algemene-voorwaarden'),
    ('WooCommerce', '/cart, /checkout, /my-account, /shop'),
    ('Gereedschap producten', '/product/pu-roller, /product/tape, /product/kwast, '
     '/product/pu-garde, /product/flexibele-spaan, /product/uitwendige-hoektroffel, '
     '/product/inwendige-hoektroffel'),
    ('Microcement producten', '/product/microcement-* (korte slugs)'),
    ('Lavasteen-gietvloeren producten', '/product/lavasteen-gietvloeren-*'),
    ('Product categorieen', '/product-category/microcement, '
     '/product-category/gereedschappen, /product-category/lavasteen-gietvloeren'),
]

table3 = doc.add_table(rows=1, cols=2, style='Light List Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Categorie'
hdr3[1].text = 'URLs'

for cat, urls in safe_categories:
    row = table3.add_row().cells
    row[0].text = cat
    row[1].text = urls

doc.add_paragraph()

# =============================================
# 6. Pagina-wijzigingen
# =============================================
doc.add_heading('6. Pagina-wijzigingen in WordPress', level=1)
doc.add_paragraph('Twee pagina\'s zijn bewerkt op 10 april:')

table4 = doc.add_table(rows=1, cols=4, style='Light List Accent 1')
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Pagina'
hdr4[1].text = 'ID'
hdr4[2].text = 'Slug'
hdr4[3].text = 'Tijdstip'

pages = [
    ('Home', '5685', '/microcement', '15:58'),
    ('Tips and Tricks', '6580', '/tips-and-tricks', '16:13'),
]
for name, pid, slug, time in pages:
    row = table4.add_row().cells
    row[0].text = name
    row[1].text = pid
    row[2].text = slug
    row[3].text = time

doc.add_paragraph()

# =============================================
# 7. Beton-cire producten op private
# =============================================
doc.add_heading('7. Beton-cire producten op private gezet', level=1)
doc.add_paragraph(
    'Alle oude beton-cire kleurproducten (Pearl White, Pure, Smooth Grey, Dark Shades, '
    'Sandy Beach, Coconut Grove, etc.) staan op "private" status. Ze zijn niet zichtbaar '
    'voor bezoekers of zoekmachines. In totaal gaat het om circa 40 producten.'
)

doc.add_paragraph()

# =============================================
# 8. Verdwenen blogposts
# =============================================
doc.add_heading('8. Verdwenen blogposts smartdeco.nl', level=1)

doc.add_heading('Situatie', level=2)
doc.add_paragraph(
    'Alle blog posts van smartdeco.nl zijn op een eerder moment verwijderd uit de database. '
    'Ze staan niet in de prullenbak of op draft -- ze zijn volledig weg. '
    'Er is geen SQL-backup beschikbaar van de smartdeco database (kpth_ prefix). '
    'De enige backups op de server zijn van epoxystone (sl14_) en betoncirewebshop (OTBgD_).'
)

doc.add_heading('Moeten de blogs terug?', level=2)
doc.add_paragraph(
    'Nee. De oude smartdeco blogs gingen over beton-cire content die niet meegaat naar '
    'de nieuwe site. Epoxystone heeft zijn eigen 12 blog posts die al live staan op '
    'epoxystone-gietvloer.nl en straks meemigreren naar smartdeco.nl. '
    'Deze blogs zijn relevanter en actueler dan de oude smartdeco content.'
)

doc.add_heading('Epoxystone blogs die meekomen na migratie', level=2)

table_blogs = doc.add_table(rows=1, cols=2, style='Light List Accent 1')
table_blogs.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_blogs = table_blogs.rows[0].cells
hdr_blogs[0].text = 'Blog titel'
hdr_blogs[1].text = 'URL slug'

epoxystone_blogs = [
    ('Microcement vloeren', '/microcement-vloeren'),
    ('Microcement keuken', '/microcement-keuken'),
    ('Microcement trap', '/microcement-trap'),
    ('Microcement meubels', '/microcement-meubels'),
    ('Microcement wanden, badkamer of toilet', '/microcement-wanden-badkamer-toilet'),
    ('Lavasteen-Gietvloeren vloeren', '/lavasteen-gietvloeren-vloeren'),
    ('Lavasteen-Gietvloeren wanden', '/lavasteen-gietvloeren-wanden'),
    ('Lavasteen-Gietvloeren douche vloer', '/lavasteen-gietvloeren-douche-vloer'),
    ('Waterdicht maken', '/waterdicht-maken'),
    ('Ondergrond opbouwen microcement of lavasteen-gietvloeren', '/ondergrond-opbouw'),
    ('Wat is een gietvloer?', '/wat-is-een-gietvloer'),
    ('Wens je een betonlook vloer of badkamer', '/wens-je-een-betonlook-vloer-of-badkamer'),
]

for title, slug in epoxystone_blogs:
    row = table_blogs.add_row().cells
    row[0].text = title
    row[1].text = slug

doc.add_paragraph()

doc.add_heading('Afhandeling oude blog-URLs', level=2)
doc.add_paragraph(
    'De oude smartdeco blog-URLs gebruikten een datum-gebaseerde structuur '
    '(bijv. /2024/02/15/blog-titel/). Deze worden al afgevangen door de 410 Gone '
    'regel in de .htaccess: ^20[0-9]{2}/[0-9]{2}/[0-9]{2}/. '
    'Google wordt dus correct geinformeerd dat deze URLs permanent verwijderd zijn.'
)
doc.add_paragraph(
    'Conclusie: geen actie nodig. De oude blogs hoeven niet hersteld te worden. '
    'De epoxystone blogs vervangen ze na de migratie.'
)

doc.add_paragraph()

# =============================================
# Samenvatting effect op Google
# =============================================
doc.add_heading('Effect op Google / SEO', level=1)

effects = [
    ('Under Construction (503)',
     'Google stopt met crawlen, behoudt huidige index tijdelijk. '
     'Na langere periode verwijdert Google de pagina\'s uit de index.'),
    ('blog_public = 0 (noindex)',
     'Alle pagina\'s krijgen noindex meta tag. Google deindexeert ze bij het volgende bezoek.'),
    ('410 Gone regels',
     'Specifieke oude URLs worden direct uit de index gehaald. '
     'Google verwerkt 410 sneller dan 404 voor deindexatie.'),
    ('WooCommerce uit',
     'Geen nieuwe bestellingen, geen productpagina\'s meer actief.'),
    ('Producten op private',
     'Zelfs als Google de 410-regels omzeilt, zijn de producten niet publiek toegankelijk.'),
]

table5 = doc.add_table(rows=1, cols=2, style='Light List Accent 1')
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr5 = table5.rows[0].cells
hdr5[0].text = 'Maatregel'
hdr5[1].text = 'Effect op Google'

for measure, effect in effects:
    row = table5.add_row().cells
    row[0].text = measure
    row[1].text = effect

for row in table5.rows:
    row.cells[0].width = Inches(2.2)
    row.cells[1].width = Inches(4.8)

doc.add_paragraph()

# --- Footer ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Document gegenereerd op 17 april 2026 - OzIS')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_dir = r'C:\Users\zeref\OneDrive\OzIS\betoncire\synvio-product-overhaul'
output_path = os.path.join(output_dir, 'Smartdeco-Migratie-Voorbereiding-10-april-2026-v2.docx')
doc.save(output_path)
print(f'Document saved: {output_path}')
